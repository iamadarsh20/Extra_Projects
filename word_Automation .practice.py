from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('''SCIENCE – BLESSING OR CURSE''',0) #Heading #0-9

p = document.add_paragraph('''Science is a great boon to human civilization.
  All signs of Progress in civilization have been made possible by science.
  Science has made our life easy and comfortable. It has given us electric fans, 
  lights. fans cool us, lights remove darkness. Lift, washing machine, etc. save 
  our labor. Car, train, bus, aircraft have made our travel speedy and comfortable.
  The computer has taken the excess load of our brain. Science has given us life-saving
  medicine. Surgery can do something miraculous. Thus science is a blessing to us. But 
  it is a curse at the same time. Science has given us speed but has taken away our emotions.
  It has made our machine. The introduction of the mobile phone has destroyed the art of letter
  writing. Science has made war more dreadful by inventing sophisticated weapons. Peace has become
  scarce. But who is responsible for making Science a curse?  ''', style='Intense Quote') #Add para
p.add_run('''Certainly, it is the evil intention of a few 
    scientists and malignant politicians. ''').bold = True
p.add_run(' .Today is Thursday')
p.add_run(' Adarsh Singh').italic = True


# document.add_paragraph(
#     'first item in unordered list', style = 'List Bullet'
# )
# document.add_paragraph(
#     'first item in unordered list', style = 'List Number'
# )

document.add_picture(r"C:\Users\kshat\Pictures\Camera Roll\WIN_20220630_14_36_34_Pro.jpg", width=Inches(1.25))



# records =(
#     (3 , '101','Spam'),
#     (7,'422','Eggs'),
#     (4,'631','Spam,spam,eggs, and spam')
# )

# for qty, id, desc in records:
#     print(qty, id, desc)
#     print('888888888888888888')






# table = document.add_table(rows =1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc


document.save('word_task.docx')

input('stop')